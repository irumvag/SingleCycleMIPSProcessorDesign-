"""
Combine Part I and Part II reports into one submission document.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer
import copy

PART1 = r'C:\Users\Chairman\Documents\1CST_UR Room\zero\Quartus2\MIPS_Project\MIPS_Project_Report_CE80665.docx'
PART2 = r'C:\Users\Chairman\Documents\1CST_UR Room\zero\Quartus2\MIPS_Project\MIPS_Project_Report_Part2_CE80665.docx'
OUT   = r'C:\Users\Chairman\Documents\1CST_UR Room\zero\Quartus2\MIPS_Project\MIPS_Project_FINAL_CE80665.docx'

# ── helpers ───────────────────────────────────────────────────────────────────
def add_page_number_field(run):
    for tag, text in [('begin', None), ('instrText', ' PAGE '), ('end', None)]:
        if tag == 'instrText':
            el = OxmlElement('w:instrText')
            el.set(qn('xml:space'), 'preserve')
            el.text = text
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), tag)
        run._r.append(el)

def shading(elem, fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    elem.append(shd)

def set_header_footer(section, part_label):
    # ---- HEADER ----
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    pPr = hp._p.get_or_add_pPr()
    # tab stop at right margin
    tabs = OxmlElement('w:tabs')
    tab  = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right'); tab.set(qn('w:pos'), '8640')
    tabs.append(tab); pPr.append(tabs)
    # bottom border
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), '1F4E79')
    pBdr.append(bot); pPr.append(pBdr)
    r1 = hp.add_run(f'CE80665 Computer Architecture — Lab Report {part_label}')
    r1.font.size = Pt(9); r1.bold = True
    r1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    hp.add_run('\t')
    r2 = hp.add_run('Group: Gad, Sandrine, Billy')
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ---- FOOTER ----
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    pPr2 = fp._p.get_or_add_pPr()
    tabs2 = OxmlElement('w:tabs')
    tab2  = OxmlElement('w:tab')
    tab2.set(qn('w:val'), 'right'); tab2.set(qn('w:pos'), '8640')
    tabs2.append(tab2); pPr2.append(tabs2)
    pBdr2 = OxmlElement('w:pBdr')
    top   = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '1');    top.set(qn('w:color'), '1F4E79')
    pBdr2.append(top); pPr2.append(pBdr2)
    r3 = fp.add_run('University of Rwanda — College of Science and Technology')
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    fp.add_run('\t')
    r4 = fp.add_run('Page '); r4.font.size = Pt(9); r4.bold = True
    r5 = fp.add_run(); r5.font.size = Pt(9); r5.bold = True
    add_page_number_field(r5)

# ── Build combined document ───────────────────────────────────────────────────
print('Opening Part I ...')
doc1 = Document(PART1)

# Fix Part I section header/footer
set_header_footer(doc1.sections[0], 'Part I')
doc1.sections[0].left_margin   = Inches(1.25)
doc1.sections[0].right_margin  = Inches(1.0)
doc1.sections[0].top_margin    = Inches(1.0)
doc1.sections[0].bottom_margin = Inches(1.0)

print('Adding Part II divider page ...')

# Add a blank "Part II" divider page at the end of Part I content
doc1.add_page_break()

# Divider heading
p = doc1.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
p.paragraph_format.space_after  = Pt(12)
r = p.add_run('PART II')
r.bold = True; r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p2 = doc1.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(8)
r2 = p2.add_run('Single-Cycle MIPS Processor')
r2.bold = True; r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p3 = doc1.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Design and Implementation in SystemVerilog on Cyclone IV E FPGA')
r3.font.size = Pt(12)

# Add section break (next page) before Part II content
# This lets us set a different header for Part II
last_para = doc1.paragraphs[-1]
pPr = last_para._p.get_or_add_pPr()
sectPr = OxmlElement('w:sectPr')
pgSz = OxmlElement('w:pgSz')
pgSz.set(qn('w:w'), '12240'); pgSz.set(qn('w:h'), '15840')
sectPr.append(pgSz)
pgMar = OxmlElement('w:pgMar')
pgMar.set(qn('w:top'),    '1440')
pgMar.set(qn('w:right'),  '1440')
pgMar.set(qn('w:bottom'), '1440')
pgMar.set(qn('w:left'),   '1800')
pgMar.set(qn('w:header'), '720')
pgMar.set(qn('w:footer'), '720')
sectPr.append(pgMar)
pPr.append(sectPr)

print('Composing Part II ...')
doc2 = Document(PART2)

composer = Composer(doc1)
composer.append(doc2)

# Fix Part II section header/footer (last section)
final_doc = doc1
if len(final_doc.sections) > 1:
    set_header_footer(final_doc.sections[-1], 'Part II')
    final_doc.sections[-1].left_margin   = Inches(1.25)
    final_doc.sections[-1].right_margin  = Inches(1.0)
    final_doc.sections[-1].top_margin    = Inches(1.0)
    final_doc.sections[-1].bottom_margin = Inches(1.0)

print(f'Saving to {OUT} ...')
composer.save(OUT)
print('Done.')
